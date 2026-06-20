from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "manual-usuario"
CAPTURES = OUT / "capturas"
DOC_CODE = "MU-FILMATE-USR-001"
VERSION = "1.0"
BASELINE_DATE = "20 de junio de 2026"
TEACHER = "Reyes Huaman, Anita Marlene"


@dataclass
class Figure:
    filename: str
    caption: str
    details: list[str]


class Manual:
    def __init__(self) -> None:
        self.md: list[str] = []
        self.doc = Document()
        self.pdf_story: list = []
        self.figure_number = 0
        self.first_level_one_heading = True
        self._configure_docx()
        self._configure_pdf()

    def _configure_docx(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.2)

        styles = self.doc.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10.5)
        styles["Normal"].paragraph_format.space_after = Pt(6)
        styles["Normal"].paragraph_format.line_spacing = 1.15

        heading_colors = {
            "Title": RGBColor(15, 23, 42),
            "Heading 1": RGBColor(30, 64, 175),
            "Heading 2": RGBColor(190, 24, 93),
            "Heading 3": RGBColor(15, 118, 110),
        }
        for name, color in heading_colors.items():
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.color.rgb = color

        styles["Heading 1"].font.size = Pt(18)
        styles["Heading 1"].paragraph_format.space_before = Pt(14)
        styles["Heading 1"].paragraph_format.space_after = Pt(8)
        styles["Heading 2"].font.size = Pt(14)
        styles["Heading 3"].font.size = Pt(11.5)

        for section in self.doc.sections:
            header = section.header.paragraphs[0]
            header.text = f"FILMATE | Manual de usuario | {DOC_CODE} | v{VERSION}"
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.runs[0].font.size = Pt(8)
            header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

            footer = section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.add_run("Página ")
            run.font.size = Pt(8)
            fld_char1 = OxmlElement("w:fldChar")
            fld_char1.set(qn("w:fldCharType"), "begin")
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = "PAGE"
            fld_char2 = OxmlElement("w:fldChar")
            fld_char2.set(qn("w:fldCharType"), "end")
            run._r.append(fld_char1)
            run._r.append(instr_text)
            run._r.append(fld_char2)
            run.add_text(" de ")
            fld_char3 = OxmlElement("w:fldChar")
            fld_char3.set(qn("w:fldCharType"), "begin")
            instr_text2 = OxmlElement("w:instrText")
            instr_text2.set(qn("xml:space"), "preserve")
            instr_text2.text = "NUMPAGES"
            fld_char4 = OxmlElement("w:fldChar")
            fld_char4.set(qn("w:fldCharType"), "end")
            run._r.append(fld_char3)
            run._r.append(instr_text2)
            run._r.append(fld_char4)

    def _configure_pdf(self) -> None:
        styles = getSampleStyleSheet()
        self.pdf_styles = {
            "title": ParagraphStyle(
                "ManualTitle",
                parent=styles["Title"],
                fontName="Helvetica-Bold",
                fontSize=23,
                leading=28,
                textColor=colors.HexColor("#0f172a"),
                alignment=TA_CENTER,
                spaceAfter=16,
            ),
            "subtitle": ParagraphStyle(
                "ManualSubtitle",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=12,
                leading=16,
                textColor=colors.HexColor("#475569"),
                alignment=TA_CENTER,
                spaceAfter=8,
            ),
            "h1": ParagraphStyle(
                "ManualH1",
                parent=styles["Heading1"],
                fontName="Helvetica-Bold",
                fontSize=17,
                leading=21,
                textColor=colors.HexColor("#1e40af"),
                spaceBefore=10,
                spaceAfter=8,
            ),
            "h2": ParagraphStyle(
                "ManualH2",
                parent=styles["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=13,
                leading=17,
                textColor=colors.HexColor("#be185d"),
                spaceBefore=9,
                spaceAfter=6,
            ),
            "h3": ParagraphStyle(
                "ManualH3",
                parent=styles["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=11,
                leading=14,
                textColor=colors.HexColor("#0f766e"),
                spaceBefore=7,
                spaceAfter=4,
            ),
            "body": ParagraphStyle(
                "ManualBody",
                parent=styles["BodyText"],
                fontName="Helvetica",
                fontSize=9.5,
                leading=13.2,
                alignment=TA_JUSTIFY,
                textColor=colors.HexColor("#1e293b"),
                spaceAfter=5,
            ),
            "caption": ParagraphStyle(
                "ManualCaption",
                parent=styles["BodyText"],
                fontName="Helvetica-Oblique",
                fontSize=8.5,
                leading=11,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#475569"),
                spaceAfter=6,
            ),
            "small": ParagraphStyle(
                "ManualSmall",
                parent=styles["BodyText"],
                fontName="Helvetica",
                fontSize=8,
                leading=10,
                textColor=colors.HexColor("#475569"),
            ),
            "callout": ParagraphStyle(
                "ManualCallout",
                parent=styles["BodyText"],
                fontName="Helvetica-Bold",
                fontSize=9.2,
                leading=12.5,
                borderColor=colors.HexColor("#93c5fd"),
                borderWidth=0.8,
                borderPadding=7,
                backColor=colors.HexColor("#eff6ff"),
                textColor=colors.HexColor("#1e3a8a"),
                spaceBefore=5,
                spaceAfter=7,
            ),
        }

    def cover(self) -> None:
        self.md += [
            "# MANUAL DE USUARIO — FILMATE",
            "",
            "**Plataforma web de cartelera, reservas, dulcería y comunidad cinematográfica**",
            "",
            "| Campo | Información |",
            "|---|---|",
            f"| Código del documento | {DOC_CODE} |",
            f"| Versión | {VERSION} |",
            f"| Línea base documental | {BASELINE_DATE} |",
            "| Curso | Gestión de la Configuración de Software |",
            "| Tipo de entregable | Manual de usuario con evidencia visual |",
            "| Estado | Versión para evaluación académica |",
            "| Elaborado por | Equipo del proyecto FILMATE |",
            f"| Docente | {TEACHER} |",
            "",
            "> Este documento describe la versión observada del frontend de usuario FILMATE y sus flujos integrados.",
            "",
        ]

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("MANUAL DE USUARIO").bold = True
        title.runs[0].font.size = Pt(27)
        title.runs[0].font.color.rgb = RGBColor(15, 23, 42)

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("FILMATE")
        run.bold = True
        run.font.size = Pt(35)
        run.font.color.rgb = RGBColor(239, 68, 68)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Plataforma web de cartelera, reservas, dulcería y comunidad cinematográfica").italic = True

        logo = ROOT / "public" / "logoGrandeFilmate.png"
        if logo.exists():
            pic = self.doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(str(logo), width=Inches(2.3))

        self.doc.add_paragraph("")
        self._doc_table(
            [
                ["Campo", "Información"],
                ["Código", DOC_CODE],
                ["Versión", VERSION],
                ["Línea base", BASELINE_DATE],
                ["Curso", "Gestión de la Configuración de Software"],
                ["Estado", "Versión para evaluación académica"],
                ["Elaborado por", "Equipo del proyecto FILMATE"],
                ["Docente", TEACHER],
            ],
            header=True,
        )
        self.doc.add_page_break()

        self.pdf_story += [
            Spacer(1, 1.5 * cm),
            Paragraph("MANUAL DE USUARIO", self.pdf_styles["title"]),
            Paragraph("FILMATE", ParagraphStyle(
                "FilmateTitle",
                parent=self.pdf_styles["title"],
                fontSize=34,
                leading=38,
                textColor=colors.HexColor("#ef4444"),
            )),
            Paragraph(
                "Plataforma web de cartelera, reservas, dulcería y comunidad cinematográfica",
                self.pdf_styles["subtitle"],
            ),
            Spacer(1, 1.2 * cm),
        ]
        if logo.exists():
            img = self._pdf_image(logo, max_w=5.3 * cm, max_h=5.3 * cm)
            self.pdf_story += [img, Spacer(1, 0.8 * cm)]
        self.pdf_story.append(
            self._pdf_table(
                [
                    ["Campo", "Información"],
                    ["Código", DOC_CODE],
                    ["Versión", VERSION],
                    ["Línea base", BASELINE_DATE],
                    ["Curso", "Gestión de la Configuración de Software"],
                    ["Estado", "Versión para evaluación académica"],
                    ["Elaborado por", "Equipo del proyecto FILMATE"],
                    ["Docente", TEACHER],
                ],
                widths=[4.7 * cm, 10.3 * cm],
            )
        )
        self.pdf_story.append(PageBreak())

    def heading(self, text: str, level: int = 1) -> None:
        should_break_before = level == 1 and not self.first_level_one_heading

        if level == 1:
            if self.first_level_one_heading:
                self.first_level_one_heading = False
            else:
                self.md += ["<div style=\"page-break-before: always;\"></div>", ""]
                self.pdf_story.append(PageBreak())

        self.md += [f"{'#' * level} {text}", ""]
        heading = self.doc.add_heading(text, level=level)
        if should_break_before:
            heading.paragraph_format.page_break_before = True
        self.pdf_story.append(Paragraph(text, self.pdf_styles[f"h{level}"]))

    def paragraph(self, text: str, callout: bool = False) -> None:
        prefix = "> " if callout else ""
        self.md += [f"{prefix}{text}", ""]
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        if callout:
            run.bold = True
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EFF6FF")
            p._p.get_or_add_pPr().append(shading)
        self.pdf_story.append(Paragraph(self._escape(text), self.pdf_styles["callout" if callout else "body"]))

    def bullets(self, items: Iterable[str], ordered: bool = False) -> None:
        items = list(items)
        for index, item in enumerate(items, start=1):
            marker = f"{index}." if ordered else "-"
            self.md.append(f"{marker} {item}")
            style = "List Number" if ordered else "List Bullet"
            p = self.doc.add_paragraph(style=style)
            p.add_run(item)
        self.md.append("")

        bullet_type = "1" if ordered else "bullet"
        self.pdf_story.append(
            ListFlowable(
                [
                    ListItem(
                        Paragraph(self._escape(item), self.pdf_styles["body"]),
                        leftIndent=12,
                    )
                    for item in items
                ],
                bulletType=bullet_type,
                leftIndent=22,
                bulletFontName="Helvetica",
                bulletFontSize=8,
                spaceAfter=6,
            )
        )

    def table(self, rows: list[list[str]]) -> None:
        self.md.append("| " + " | ".join(rows[0]) + " |")
        self.md.append("|" + "|".join(["---"] * len(rows[0])) + "|")
        for row in rows[1:]:
            self.md.append("| " + " | ".join(row) + " |")
        self.md.append("")
        self._doc_table(rows, header=True)
        col_count = len(rows[0])
        usable = 15.2 * cm
        widths = [usable / col_count] * col_count
        self.pdf_story += [self._pdf_table(rows, widths=widths), Spacer(1, 0.2 * cm)]

    def figure(self, figure: Figure) -> None:
        self.figure_number += 1
        path = CAPTURES / figure.filename
        caption = f"Figura {self.figure_number}. {figure.caption}"
        rel = Path("capturas") / figure.filename
        self.md += [f"![{caption}]({rel.as_posix()})", "", f"*{caption}*", ""]
        if figure.details:
            self.md.append("Elementos y lectura de la pantalla:")
            self.md.append("")
            for item in figure.details:
                self.md.append(f"- {item}")
            self.md.append("")

        if path.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width, height = self._doc_image_size(path)
            p.add_run().add_picture(str(path), width=width, height=height)
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(caption)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(71, 85, 105)
            for item in figure.details:
                bp = self.doc.add_paragraph(style="List Bullet")
                bp.add_run(item)

            pdf_img = self._pdf_image(path, max_w=17.1 * cm, max_h=17.2 * cm)
            block = [pdf_img, Spacer(1, 0.12 * cm), Paragraph(caption, self.pdf_styles["caption"])]
            self.pdf_story += [KeepTogether(block)]
            if figure.details:
                self.pdf_story.append(
                    ListFlowable(
                        [
                            ListItem(Paragraph(self._escape(item), self.pdf_styles["small"]), leftIndent=10)
                            for item in figure.details
                        ],
                        bulletType="bullet",
                        leftIndent=20,
                        bulletFontSize=7,
                        spaceAfter=6,
                    )
                )

    def page_break(self) -> None:
        self.md += ["<div style=\"page-break-after: always;\"></div>", ""]
        self.doc.add_page_break()
        self.pdf_story.append(PageBreak())

    def save(self) -> None:
        OUT.mkdir(parents=True, exist_ok=True)
        md_path = OUT / "Manual_de_Usuario_FILMATE.md"
        docx_path = OUT / "Manual_de_Usuario_FILMATE.docx"
        pdf_path = OUT / "Manual_de_Usuario_FILMATE.pdf"

        md_path.write_text("\n".join(self.md), encoding="utf-8")
        self.doc.save(docx_path)

        pdf = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            rightMargin=1.7 * cm,
            leftMargin=1.7 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.7 * cm,
            title="Manual de Usuario FILMATE",
            author="Equipo FILMATE",
            subject="Gestión de la Configuración de Software",
        )
        pdf.build(self.pdf_story, onFirstPage=self._pdf_page, onLaterPages=self._pdf_page)

        print(md_path)
        print(docx_path)
        print(pdf_path)

    def _doc_table(self, rows: list[list[str]], header: bool = True) -> None:
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for r_index, row in enumerate(rows):
            for c_index, value in enumerate(row):
                cell = table.cell(r_index, c_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if header and r_index == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                if header and r_index == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "1E3A8A")
                    cell._tc.get_or_add_tcPr().append(shading)
        self.doc.add_paragraph("")

    def _pdf_table(self, rows: list[list[str]], widths: list[float]) -> Table:
        data = []
        for row_index, row in enumerate(rows):
            style = ParagraphStyle(
                f"TableCell{row_index}",
                parent=self.pdf_styles["small"],
                textColor=colors.white if row_index == 0 else colors.HexColor("#1e293b"),
                fontName="Helvetica-Bold" if row_index == 0 else "Helvetica",
                leading=10,
            )
            data.append([Paragraph(self._escape(str(cell)), style) for cell in row])
        table = Table(data, colWidths=widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]
            )
        )
        return table

    def _doc_image_size(self, path: Path) -> tuple:
        with PILImage.open(path) as image:
            px_w, px_h = image.size
        max_w = 6.45
        max_h = 7.1
        scale = min(max_w / px_w, max_h / px_h)
        return Inches(px_w * scale), Inches(px_h * scale)

    def _pdf_image(self, path: Path, max_w: float, max_h: float) -> Image:
        with PILImage.open(path) as image:
            px_w, px_h = image.size
        scale = min(max_w / px_w, max_h / px_h)
        return Image(str(path), width=px_w * scale, height=px_h * scale)

    def _pdf_page(self, canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(1.7 * cm, A4[1] - 1.05 * cm, f"FILMATE | {DOC_CODE} | v{VERSION}")
        canvas.drawRightString(A4[0] - 1.7 * cm, 0.9 * cm, f"Página {doc.page}")
        canvas.restoreState()

    @staticmethod
    def _escape(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )


def add_control_and_scope(manual: Manual) -> None:
    manual.heading("Control del documento", 1)
    manual.table(
        [
            ["Versión", "Fecha", "Descripción del cambio", "Responsable"],
            ["0.1", "20/06/2026", "Inventario funcional y planificación de capturas.", "Equipo FILMATE"],
            ["0.9", "20/06/2026", "Ejecución de 24 evidencias visuales y validación de flujos.", "Equipo FILMATE"],
            ["1.0", "20/06/2026", "Emisión del manual académico en Markdown, Word y PDF.", "Equipo FILMATE"],
        ]
    )
    manual.paragraph(
        "Criterio de configuración documental: este manual se identifica mediante el código "
        f"{DOC_CODE}; la versión {VERSION} queda asociada a la línea base visual del {BASELINE_DATE}. "
        "Cualquier cambio en rutas, etiquetas, reglas de validación, pasos de compra o diseño de pantallas "
        "debe originar una nueva revisión del documento y la sustitución de las capturas afectadas.",
        callout=True,
    )

    manual.heading("Contenido", 1)
    manual.bullets(
        [
            "1. Introducción, objetivo y alcance.",
            "2. Descripción del sistema, perfiles y requisitos de uso.",
            "3. Convenciones visuales y mapa de navegación.",
            "4. Acceso, registro e inicio de sesión.",
            "5. Cartelera, filtros y detalle de película.",
            "6. Cines y ubicación de locales.",
            "7. Reserva de función y selección de asientos.",
            "8. Dulcería, pago, comprobante y QR.",
            "9. Perfil social, búsqueda, seguimiento y edición.",
            "10. Cierre de sesión, solución de problemas y seguridad.",
            "11. Trazabilidad, criterios de aceptación y gestión del manual.",
            "12. Glosario y guía rápida.",
        ]
    )

    manual.heading("1. Introducción", 1)
    manual.heading("1.1 Objetivo", 2)
    manual.paragraph(
        "Orientar al usuario final en el uso seguro y correcto de FILMATE, desde el acceso inicial hasta la "
        "obtención del ticket de compra y la administración del perfil social. Cada procedimiento incluye "
        "precondiciones, acciones, resultado esperado y evidencia visual."
    )
    manual.heading("1.2 Alcance", 2)
    manual.paragraph(
        "El manual cubre el frontend de usuario: inicio de sesión, modo invitado, registro, cartelera, filtros, "
        "detalle de películas, horarios, cines, selección de asientos, dulcería, verificación, pago, ticket QR, "
        "perfil social, búsqueda de usuarios, seguimiento, edición de perfil, favoritos y cierre de sesión."
    )
    manual.paragraph(
        "No cubre la administración interna de películas, cines, salas, funciones, usuarios o transacciones; "
        "esas operaciones pertenecen al frontend administrativo y requieren un manual independiente."
    )
    manual.heading("1.3 Base de la evidencia", 2)
    manual.paragraph(
        "Las capturas se obtuvieron el 20 de junio de 2026 sobre la aplicación React/Vite del repositorio "
        "FILMATE_UserFrontend. Debido a que la carga SQL de referencia programaba funciones únicamente del "
        "6 al 16 de junio de 2026, se utilizó un entorno de demostración local controlado con la misma estructura "
        "de datos y contratos de API para habilitar fechas vigentes, asientos, dulcería, compra y módulos sociales. "
        "No se utilizaron datos financieros reales."
    )

    manual.heading("2. Descripción general del sistema", 1)
    manual.paragraph(
        "FILMATE es una plataforma web orientada a clientes de cine. Integra consulta de cartelera, reserva de "
        "butacas, compra de productos de dulcería, emisión de comprobante con QR y una comunidad social de gustos cinematográficos."
    )
    manual.heading("2.1 Perfiles de acceso", 2)
    manual.table(
        [
            ["Perfil", "Capacidades principales", "Restricciones"],
            ["Invitado", "Consultar cartelera, detalle de películas, cines y dulcería.", "No accede al módulo Social ni completa reservas asociadas a usuario."],
            ["Usuario registrado", "Todas las funciones públicas, reserva, pago, ticket y perfil social.", "Debe mantener una sesión válida y datos de contacto correctos."],
        ]
    )
    manual.heading("2.2 Requisitos mínimos", 2)
    manual.bullets(
        [
            "Equipo de escritorio, portátil, tableta o teléfono con navegador moderno.",
            "Google Chrome, Microsoft Edge o Mozilla Firefox actualizado.",
            "Resolución recomendada de escritorio: 1366 × 768 o superior.",
            "Conexión de red estable con acceso al frontend y al servicio API.",
            "Correo electrónico válido para usuarios registrados.",
            "JavaScript y almacenamiento local habilitados en el navegador.",
            "Para descargar el ticket: permisos de descarga habilitados.",
        ]
    )
    manual.heading("2.3 Datos solicitados al registrarse", 2)
    manual.table(
        [
            ["Campo", "Obligatorio", "Regla aplicada"],
            ["Nombre completo", "Sí", "Solo letras y espacios."],
            ["Nombre de usuario", "Sí", "Entre 3 y 20 caracteres; letras, números y guion bajo."],
            ["Correo electrónico", "Sí", "Formato usuario@dominio."],
            ["Contraseña", "Sí", "Mínimo 6 caracteres."],
            ["Documento", "Sí", "Entre 8 y 15 caracteres alfanuméricos."],
            ["Teléfono", "No", "Entre 7 y 15 dígitos, sin letras."],
        ]
    )

    manual.heading("3. Convenciones y navegación", 1)
    manual.heading("3.1 Convenciones visuales", 2)
    manual.table(
        [
            ["Elemento", "Significado"],
            ["Botón rojo", "Acción principal sensible o cierre de sesión."],
            ["Botón azul", "Acción de navegación, edición o selección."],
            ["Botón verde", "Confirmación o avance de compra."],
            ["Mensaje rojo", "Error de validación o imposibilidad de completar una operación."],
            ["Mensaje verde", "Operación completada correctamente."],
            ["Asiento claro", "Disponible para selección."],
            ["Asiento verde", "Seleccionado por el usuario."],
            ["Asiento rojo", "Ocupado o no disponible."],
        ]
    )
    manual.heading("3.2 Mapa de navegación", 2)
    manual.paragraph(
        "La cabecera permite desplazarse entre Cartelera, Cines, Dulcería y Social. El módulo Social aparece "
        "únicamente cuando la sesión corresponde a un usuario registrado. La ruta de compra sigue la secuencia:"
    )
    manual.paragraph(
        "Cartelera → Detalle de película → Función → Asientos → Dulcería → Verificación → Pago → Ticket.",
        callout=True,
    )


def add_access_and_registration(manual: Manual) -> None:
    manual.heading("4. Acceso al sistema", 1)
    manual.heading("4.1 Iniciar sesión con cuenta registrada", 2)
    manual.paragraph("Precondición: el usuario debe disponer de una cuenta activa.")
    manual.bullets(
        [
            "Abra la URL de FILMATE en el navegador.",
            "Escriba el correo electrónico registrado.",
            "Escriba la contraseña; el valor se oculta por seguridad.",
            "Seleccione «Iniciar sesión» o presione Enter.",
            "Espere el mensaje de éxito y la redirección automática a la cartelera.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "01-inicio-sesion.png",
            "Pantalla de inicio de sesión.",
            [
                "El campo de correo exige un formato que incluya el carácter @.",
                "La contraseña se muestra en modo oculto.",
                "El usuario puede ir al registro o entrar como invitado.",
            ],
        )
    )

    manual.heading("4.2 Validaciones de acceso", 2)
    manual.paragraph(
        "Si falta un dato, el correo no tiene formato válido o las credenciales no coinciden, FILMATE conserva al "
        "usuario en la pantalla y presenta un mensaje. Corrija el dato indicado y vuelva a intentar; no es necesario recargar la página."
    )
    manual.figure(
        Figure(
            "02-validacion-inicio-sesion.png",
            "Ejemplo de validación de correo.",
            [
                "El mensaje aparece encima de los campos.",
                "La validación evita enviar un correo sin formato válido al backend.",
            ],
        )
    )

    manual.heading("4.3 Entrar como invitado", 2)
    manual.bullets(
        [
            "En la pantalla de acceso, seleccione «entrar como invitado».",
            "Espere la confirmación «Bienvenido, Invitado».",
            "Consulte cartelera, cines y dulcería.",
            "Para reservar o usar Social, cierre la sesión de invitado e ingrese con una cuenta registrada.",
        ],
        ordered=True,
    )
    manual.paragraph(
        "El modo invitado es útil para exploración rápida, pero no debe emplearse para operaciones que necesiten "
        "identificar al comprador o conservar preferencias sociales.",
        callout=True,
    )

    manual.heading("5. Registro de usuario", 1)
    manual.heading("5.1 Abrir el formulario", 2)
    manual.paragraph(
        "Desde el inicio de sesión, seleccione «Regístrate». El formulario organiza los datos en dos columnas en "
        "escritorio y en una sola columna en pantallas pequeñas."
    )
    manual.figure(
        Figure(
            "03-registro-usuario.png",
            "Formulario de creación de cuenta.",
            [
                "La columna izquierda agrupa nombre, correo y contraseña.",
                "La columna derecha agrupa username, documento y teléfono.",
                "El teléfono es el único campo opcional.",
            ],
        )
    )

    manual.heading("5.2 Completar y enviar el registro", 2)
    manual.bullets(
        [
            "Ingrese el nombre completo sin números ni símbolos.",
            "Defina un nombre de usuario único de 3 a 20 caracteres.",
            "Ingrese un correo válido y una contraseña de al menos 6 caracteres.",
            "Registre el documento entre 8 y 15 caracteres.",
            "Opcionalmente, ingrese un teléfono de 7 a 15 dígitos.",
            "Revise los datos y seleccione «Registrarse».",
            "Espere el mensaje «Registro exitoso» y la redirección a la cartelera.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "04-registro-completado.png",
            "Formulario diligenciado antes del envío.",
            [
                "Revise especialmente correo, documento y username porque identifican la cuenta.",
                "No comparta la contraseña ni utilice una clave de otro servicio.",
            ],
        )
    )

    manual.heading("5.3 Errores frecuentes de registro", 2)
    manual.table(
        [
            ["Mensaje o situación", "Causa probable", "Acción recomendada"],
            ["Correo inválido", "Falta @, dominio o extensión.", "Corregir el formato del correo."],
            ["Username ya registrado", "El nombre está en uso.", "Probar una variante única."],
            ["Teléfono ya registrado", "Duplicidad en el registro local o backend.", "Verificar el número o usar otro."],
            ["Contraseña muy corta", "Menos de 6 caracteres.", "Definir una contraseña más larga."],
            ["Documento inválido", "Longitud o caracteres no admitidos.", "Usar 8 a 15 caracteres alfanuméricos."],
            ["No se pudo completar", "API o red no disponible.", "Esperar, comprobar conexión y reintentar."],
        ]
    )


def add_catalog_and_cinemas(manual: Manual) -> None:
    manual.heading("6. Cartelera", 1)
    manual.heading("6.1 Consultar recomendaciones y películas", 2)
    manual.paragraph(
        "Después del acceso, la cartelera presenta recomendaciones destacadas, filtros y el catálogo disponible. "
        "Cada tarjeta incluye póster, título, género, duración, clasificación y valoración."
    )
    manual.figure(
        Figure(
            "05-cartelera-principal.png",
            "Vista general de la cartelera.",
            [
                "La cabecera mantiene visibles las rutas principales y el cierre de sesión.",
                "Las recomendaciones aparecen primero.",
                "Los filtros preceden a la cuadrícula completa de cartelera.",
            ],
        )
    )

    manual.heading("6.2 Aplicar filtros", 2)
    manual.bullets(
        [
            "En «Filtrar cartelera», elija un día que tenga funciones.",
            "Seleccione un cine específico o mantenga «Todos los cines».",
            "Seleccione un género o mantenga «Todos los géneros».",
            "Revise el resultado actualizado debajo del panel.",
            "Use «Limpiar filtros» cuando exista al menos una selección distinta del valor general.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "06-filtros-cartelera.png",
            "Cartelera filtrada por sede.",
            [
                "Los filtros se combinan: día + cine + género.",
                "Si no hay coincidencias, cambie uno de los criterios.",
            ],
        )
    )

    manual.heading("6.3 Abrir el detalle de una película", 2)
    manual.bullets(
        [
            "Ubique la película en recomendaciones o cartelera.",
            "Seleccione su tarjeta.",
            "Revise título, géneros, clasificación, duración, sinopsis, dirección y reparto.",
            "Consulte el tráiler y las reseñas disponibles.",
            "Desplácese hasta los horarios por cine para iniciar una reserva.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "07-detalle-pelicula.png",
            "Detalle completo de película.",
            [
                "La información proviene del catálogo configurado en el backend.",
                "El botón de regreso conserva el contexto de navegación.",
                "Los horarios se agrupan por sede.",
            ],
        )
    )

    manual.heading("6.4 Consultar horarios", 2)
    manual.paragraph(
        "Cada sede muestra las funciones de la fecha activa. La hora y sala deben verificarse antes de continuar, "
        "pues determinan el precio y el inventario de butacas."
    )
    manual.figure(
        Figure(
            "08-horarios-por-cine.png",
            "Horarios disponibles por cine.",
            [
                "Seleccione exactamente la función que desea reservar.",
                "Si no aparecen horarios, pruebe otra película o fecha.",
            ],
        )
    )

    manual.heading("7. Cines", 1)
    manual.heading("7.1 Consultar sedes", 2)
    manual.bullets(
        [
            "Seleccione «Cines» en la cabecera.",
            "Revise nombre, dirección y horario de atención de cada sede.",
            "Use el mapa embebido para reconocer la ubicación.",
            "Seleccione «Ver en Google Maps» para ampliar la ubicación dentro de FILMATE.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "17-listado-cines.png",
            "Listado de locales FILMATE.",
            [
                "Cada tarjeta mantiene juntos mapa, dirección y horario.",
                "La disponibilidad de mapas externos depende de la conexión de red.",
            ],
        )
    )
    manual.figure(
        Figure(
            "18-mapa-cine.png",
            "Ampliación de la ubicación de una sede.",
            [
                "Cierre la ventana para volver al listado.",
                "Verifique la dirección textual si el mapa no carga.",
            ],
        )
    )


def add_reservation_and_checkout(manual: Manual) -> None:
    manual.heading("8. Reserva de entradas", 1)
    manual.heading("8.1 Precondiciones", 2)
    manual.bullets(
        [
            "Sesión iniciada como usuario registrado.",
            "Película, fecha, sede y función seleccionadas.",
            "Mapa de asientos cargado.",
            "Conexión activa durante la selección y el pago.",
        ]
    )

    manual.heading("8.2 Interpretar el mapa de asientos", 2)
    manual.paragraph(
        "La pantalla se ubica en la parte superior. Cada butaca muestra número y fila. El panel izquierdo resume "
        "película, cine, fecha, hora y sala. La disponibilidad puede cambiar si otro usuario reserva simultáneamente."
    )
    manual.figure(
        Figure(
            "09-mapa-asientos.png",
            "Mapa de butacas antes de seleccionar.",
            [
                "Claro: asiento disponible.",
                "Rojo: asiento ocupado o bloqueado.",
                "El botón «Siguiente» permanece inactivo mientras no haya selección.",
            ],
        )
    )

    manual.heading("8.3 Seleccionar butacas", 2)
    manual.bullets(
        [
            "Seleccione una butaca clara.",
            "Repita la acción para agregar más asientos.",
            "Compruebe que las butacas elegidas cambien a color verde.",
            "Para retirar una selección, vuelva a presionar la misma butaca.",
            "Revise el contador de seleccionados y pulse «Siguiente».",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "10-asientos-seleccionados.png",
            "Butacas D5 y D6 seleccionadas.",
            [
                "El color verde identifica temporalmente las butacas del usuario.",
                "No cierre la pestaña durante una selección activa.",
            ],
        )
    )

    manual.heading("8.4 Confirmar asientos y total", 2)
    manual.paragraph(
        "FILMATE presenta una confirmación con identificadores, cantidad y total de entradas. Use «Revisar» para "
        "volver al mapa o «Confirmar» para continuar a dulcería."
    )
    manual.figure(
        Figure(
            "11-confirmacion-asientos.png",
            "Confirmación previa de butacas.",
            [
                "El total se calcula como cantidad de asientos por precio base de la función.",
                "Después de confirmar, conserve la misma sesión hasta obtener el ticket.",
            ],
        )
    )

    manual.heading("9. Dulcería y carrito", 1)
    manual.heading("9.1 Flujo asociado a una reserva", 2)
    manual.paragraph(
        "La parte superior muestra la reserva seleccionada. El catálogo se organiza en Combos, Canchita, Bebidas y "
        "Dulces; el carrito permanece visible en escritorio."
    )
    manual.figure(
        Figure(
            "12-dulceria-reserva.png",
            "Dulcería abierta desde una reserva.",
            [
                "«Volver a asientos» permite corregir la selección antes del pago.",
                "«Omitir snacks» continúa directamente con el importe de entradas.",
            ],
        )
    )

    manual.heading("9.2 Agregar productos", 2)
    manual.bullets(
        [
            "Seleccione «Agregar» en el producto deseado.",
            "Use los botones + y − del carrito para ajustar la cantidad.",
            "Compruebe los subtotales de asientos y snacks.",
            "Use «Omitir snacks» si desea comprar solo entradas.",
            "Seleccione «Confirmar pedido» cuando el resumen sea correcto.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "13-carrito-dulceria.png",
            "Carrito con un combo agregado.",
            [
                "El total general combina entradas y productos.",
                "Una cantidad reducida a cero elimina el producto.",
            ],
        )
    )

    manual.heading("10. Verificación y pago", 1)
    manual.heading("10.1 Verificar la compra", 2)
    manual.paragraph(
        "Antes de pagar, revise productos, cantidades, subtotal de snacks, subtotal de asientos y total. "
        "Seleccione «Cancelar pedido» para corregir o «Pagar ahora» para avanzar."
    )
    manual.figure(
        Figure(
            "14-verificacion-compra.png",
            "Resumen de verificación.",
            [
                "Esta es la última revisión de cantidades antes de elegir el medio de pago.",
                "El usuario todavía puede regresar sin emitir la transacción.",
            ],
        )
    )

    manual.heading("10.2 Elegir método de pago", 2)
    manual.bullets(
        [
            "Revise el total a cobrar y los datos de película, sede, función y asientos.",
            "Seleccione Tarjeta, Yape, Plin o Efectivo.",
            "Compruebe el texto «Método seleccionado».",
            "Pulse «Pagar con …».",
            "No cierre ni actualice la página mientras aparezca «Procesando».",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "15-metodo-pago.png",
            "Selección del método de pago.",
            [
                "El medio activo se resalta visualmente.",
                "En un entorno productivo, siga las instrucciones del proveedor de pago.",
            ],
        )
    )

    manual.heading("10.3 Resultado esperado del pago", 2)
    manual.paragraph(
        "La operación correcta genera un identificador de transacción, un número de pedido, el detalle del consumo, "
        "el total y un QR. Si aparece un error, no repita el pago inmediatamente: primero compruebe si ya existe una transacción."
    )

    manual.heading("11. Ticket y comprobante QR", 1)
    manual.bullets(
        [
            "Verifique película, sede, fecha, sala y asientos.",
            "Verifique productos y total.",
            "Conserve el número de pedido y el identificador de transacción.",
            "Seleccione «Descargar PDF y salir» para guardar una copia.",
            "Presente el QR en el punto de atención correspondiente.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "16-ticket-compra.png",
            "Ticket de compra con QR.",
            [
                "El QR representa la transacción y no debe publicarse en redes sociales.",
                "La descarga genera un archivo PDF en la carpeta configurada por el navegador.",
                "El botón de regreso finaliza el flujo y vuelve a la cartelera.",
            ],
        )
    )
    manual.paragraph(
        "Recomendación de seguridad: trate el QR como un comprobante personal. Una captura compartida puede permitir "
        "que otra persona intente utilizar el pedido antes que el comprador.",
        callout=True,
    )


def add_social_and_logout(manual: Manual) -> None:
    manual.heading("12. Módulo Social", 1)
    manual.heading("12.1 Consultar el perfil propio", 2)
    manual.paragraph(
        "El perfil presenta nombre, username, avatar, estadísticas, biografía, distribución personal de "
        "calificaciones y hasta cinco películas favoritas."
    )
    manual.figure(
        Figure(
            "19-perfil-social.png",
            "Perfil social de usuario registrado.",
            [
                "Las estadísticas muestran películas, seguidos y seguidores.",
                "La clasificación personal resume las valoraciones por estrellas.",
                "«Editar Perfil» solo aparece en el perfil propio.",
            ],
        )
    )

    manual.heading("12.2 Buscar usuarios", 2)
    manual.bullets(
        [
            "Escriba al menos dos caracteres del username en «Buscar username».",
            "Espere la lista de coincidencias.",
            "Revise username y nombre del resultado.",
            "Seleccione el usuario para abrir su perfil.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "20-busqueda-usuarios.png",
            "Resultados de búsqueda social.",
            [
                "La búsqueda prioriza usernames que comienzan con el texto.",
                "Si no hay coincidencias, pruebe una parte diferente del username.",
            ],
        )
    )

    manual.heading("12.3 Seguir a otro usuario", 2)
    manual.bullets(
        [
            "Abra el perfil de otra persona desde la búsqueda.",
            "Revise que no sea su propio perfil.",
            "Seleccione «Seguir».",
            "Espere que el botón cambie a «Siguiendo» y que aumente el contador.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "21-perfil-otro-usuario.png",
            "Perfil de otro usuario con acción de seguimiento.",
            [
                "El botón «Seguir» sustituye a «Editar Perfil».",
                "La operación requiere una sesión registrada.",
            ],
        )
    )

    manual.heading("12.4 Editar información del perfil", 2)
    manual.bullets(
        [
            "En el perfil propio, seleccione «Editar Perfil».",
            "Pulse «Modificar» para habilitar nombre, username, correo, teléfono, bio y avatar.",
            "Realice los cambios necesarios.",
            "Use «Cambiar avatar» para elegir una imagen propuesta.",
            "Seleccione «Guardar cambios» y espere la confirmación.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "22-editar-perfil.png",
            "Pantalla de edición del perfil.",
            [
                "Los campos están bloqueados hasta seleccionar «Modificar».",
                "La bio se presenta públicamente en el perfil social.",
                "El cambio de correo debe realizarse con especial cuidado.",
            ],
        )
    )

    manual.heading("12.5 Administrar películas favoritas", 2)
    manual.bullets(
        [
            "En «Películas favoritas», seleccione «Editar».",
            "Busque una película por título si es necesario.",
            "Seleccione o deseleccione tarjetas; el máximo es cinco.",
            "Pulse «Aplicar selección».",
            "Finalmente, seleccione «Guardar cambios» en la pantalla de edición.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "23-seleccionar-favoritas.png",
            "Selector de películas favoritas.",
            [
                "El contador superior indica cuántas de cinco están seleccionadas.",
                "Una marca azul identifica cada favorita.",
                "Al alcanzar cinco, las películas restantes quedan deshabilitadas hasta retirar una.",
            ],
        )
    )

    manual.heading("13. Cerrar sesión", 1)
    manual.bullets(
        [
            "Seleccione «Cerrar Sesión» en la cabecera.",
            "Revise el mensaje de confirmación.",
            "Pulse «Cancelar» para continuar o «Cerrar Sesión» para finalizar.",
            "Compruebe que el sistema vuelva a la pantalla de acceso.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "24-cerrar-sesion.png",
            "Confirmación de cierre de sesión.",
            [
                "El cierre elimina la sesión almacenada en el navegador.",
                "En equipos compartidos, cierre siempre la sesión al terminar.",
            ],
        )
    )


def add_support_and_gcs(manual: Manual) -> None:
    manual.heading("14. Solución de problemas", 1)
    manual.table(
        [
            ["Problema", "Causa probable", "Solución del usuario"],
            ["No cargan películas", "Backend no disponible o red interrumpida.", "Comprobar conexión, esperar y recargar una vez."],
            ["No aparecen días con funciones", "No existe programación vigente.", "Consultar otra fecha o informar al administrador."],
            ["No aparece el mapa de asientos", "La función no tiene inventario o falló la API.", "Volver y elegir otra función; reintentar después."],
            ["Un asiento cambia a ocupado", "Otro usuario lo reservó simultáneamente.", "Seleccionar otra butaca disponible."],
            ["No se puede pagar", "Sesión de invitado, datos incompletos o error de red.", "Iniciar sesión, verificar reserva y reintentar una sola vez."],
            ["No se descarga el PDF", "Descargas bloqueadas por el navegador.", "Autorizar descargas y volver a usar el botón."],
            ["Social redirige al inicio", "No existe sesión registrada.", "Iniciar sesión con una cuenta activa."],
            ["Avatar o mapa no cargan", "Recurso externo inaccesible.", "Verificar internet; el resto de datos sigue disponible."],
            ["Texto o botones se superponen", "Zoom o resolución extrema.", "Restablecer zoom al 100 % y usar una resolución recomendada."],
        ]
    )
    manual.heading("14.1 Qué hacer ante un pago incierto", 2)
    manual.bullets(
        [
            "No presione repetidamente el botón de pago.",
            "Revise si apareció número de pedido o transacción.",
            "Conserve la hora, usuario, película, sede y total.",
            "Consulte el historial disponible o contacte soporte antes de intentar de nuevo.",
            "No comparta datos completos de tarjeta ni el QR con soporte informal.",
        ],
        ordered=True,
    )

    manual.heading("15. Seguridad, privacidad y buenas prácticas", 1)
    manual.bullets(
        [
            "Use una contraseña exclusiva y no la comparta.",
            "Compruebe la URL antes de ingresar credenciales.",
            "Cierre sesión en equipos públicos o compartidos.",
            "No publique tickets, QR, documentos, correo o teléfono.",
            "Revise el total antes de confirmar el pago.",
            "Evite actualizar o retroceder mientras se procesa una transacción.",
            "Mantenga el navegador actualizado.",
            "Reporte mensajes inesperados, cobros duplicados o cambios de perfil no autorizados.",
        ]
    )

    manual.heading("16. Criterios de aceptación del usuario", 1)
    manual.table(
        [
            ["ID", "Criterio verificable", "Evidencia"],
            ["CA-01", "El usuario puede iniciar sesión o recibir una validación comprensible.", "Figuras 1 y 2"],
            ["CA-02", "El usuario puede registrar una cuenta con reglas visibles.", "Figuras 3 y 4"],
            ["CA-03", "La cartelera lista y filtra películas.", "Figuras 5 y 6"],
            ["CA-04", "El detalle muestra información y funciones por cine.", "Figuras 7 y 8"],
            ["CA-05", "El usuario puede seleccionar y confirmar asientos.", "Figuras 9, 10 y 11"],
            ["CA-06", "La dulcería permite agregar productos y calcular totales.", "Figuras 12 y 13"],
            ["CA-07", "La compra se verifica, paga y genera ticket QR.", "Figuras 14, 15 y 16"],
            ["CA-08", "El usuario consulta sedes y mapas.", "Figuras 17 y 18"],
            ["CA-09", "El perfil social muestra estadísticas y favoritos.", "Figura 19"],
            ["CA-10", "El usuario busca y sigue otros perfiles.", "Figuras 20 y 21"],
            ["CA-11", "El usuario edita su perfil y favoritas.", "Figuras 22 y 23"],
            ["CA-12", "El usuario puede cerrar su sesión de forma confirmada.", "Figura 24"],
        ]
    )

    manual.heading("17. Gestión de configuración del manual", 1)
    manual.paragraph(
        "Para el curso de Gestión de la Configuración de Software, este manual se considera un elemento de "
        "configuración documental relacionado con el frontend, el contrato API y la línea base de datos de demostración."
    )
    manual.heading("17.1 Elementos bajo control", 2)
    manual.table(
        [
            ["Elemento", "Identificador / ubicación", "Disparador de actualización"],
            ["Manual editable", "docs/manual-usuario/Manual_de_Usuario_FILMATE.docx", "Cambio funcional o corrección aprobada."],
            ["Manual portable", "docs/manual-usuario/Manual_de_Usuario_FILMATE.pdf", "Nueva versión del manual editable."],
            ["Fuente trazable", "docs/manual-usuario/Manual_de_Usuario_FILMATE.md", "Cambio de contenido o estructura."],
            ["Evidencias", "docs/manual-usuario/capturas/*.png", "Cambio visual, etiqueta, ruta o resultado."],
            ["Automatización", "scripts/manual/capture-screenshots.mjs", "Cambio de selectores o flujo."],
            ["Entorno de evidencia", "scripts/manual/mock-server.mjs", "Cambio de contrato de API o datos necesarios."],
        ]
    )
    manual.heading("17.2 Procedimiento de cambio", 2)
    manual.bullets(
        [
            "Registrar la solicitud de cambio y el motivo.",
            "Identificar requisitos, pantallas y procedimientos afectados.",
            "Actualizar la aplicación o el contrato API correspondiente.",
            "Ejecutar nuevamente las capturas afectadas.",
            "Actualizar contenido, control de versiones y matriz de trazabilidad.",
            "Revisar ortografía, consistencia, enlaces e imágenes.",
            "Generar Word y PDF desde la fuente.",
            "Aprobar y etiquetar la nueva línea base documental.",
        ],
        ordered=True,
    )
    manual.heading("17.3 Nomenclatura de versiones", 2)
    manual.paragraph(
        "Se recomienda usar versión mayor cuando cambia el flujo o alcance del manual; versión menor cuando se "
        "agregan funciones compatibles; y revisión de parche cuando solo se corrigen redacción, formato o capturas sin cambiar el procedimiento."
    )

    manual.heading("18. Glosario", 1)
    manual.table(
        [
            ["Término", "Definición"],
            ["API", "Servicio que comunica el frontend con datos y operaciones del sistema."],
            ["Cartelera", "Conjunto de películas con funciones disponibles."],
            ["Función", "Exhibición de una película en fecha, hora, cine y sala determinados."],
            ["Butaca / asiento", "Lugar individual que puede reservarse para una función."],
            ["Bloqueo temporal", "Reserva breve de un asiento mientras el usuario completa la compra."],
            ["Dulcería", "Catálogo de combos, canchita, bebidas y dulces."],
            ["QR", "Código gráfico asociado al ticket o transacción."],
            ["Username", "Nombre público y único utilizado en el módulo Social."],
            ["Línea base", "Versión aprobada y controlada de software, datos o documentación."],
            ["Elemento de configuración", "Activo sujeto a identificación, versión, cambio y auditoría."],
        ]
    )

    manual.heading("19. Guía rápida", 1)
    manual.paragraph(
        "Comprar entradas: Iniciar sesión → Cartelera → Película → Horario → Asientos → Confirmar → "
        "Agregar u omitir snacks → Verificar → Elegir pago → Pagar → Descargar ticket.",
        callout=True,
    )
    manual.paragraph(
        "Actualizar perfil: Social → Editar Perfil → Modificar → Cambiar datos/avatar → Editar favoritas → "
        "Aplicar selección → Guardar cambios.",
        callout=True,
    )
    manual.heading("20. Verificación de integridad del entregable", 1)
    manual.bullets(
        [
            "La portada identifica el curso, la docente, la versión y la línea base documental.",
            "El PDF contiene las 24 figuras correspondientes a los flujos documentados.",
            "El archivo DOCX conserva el contenido editable y la paginación.",
            "El documento no incluye contraseñas reales ni datos financieros sensibles.",
            "Los archivos fuente y las evidencias visuales permanecen versionados en el repositorio.",
        ]
    )


def build_manual() -> None:
    manual = Manual()
    manual.cover()
    add_control_and_scope(manual)
    add_access_and_registration(manual)
    add_catalog_and_cinemas(manual)
    add_reservation_and_checkout(manual)
    add_social_and_logout(manual)
    add_support_and_gcs(manual)
    manual.save()


if __name__ == "__main__":
    build_manual()
